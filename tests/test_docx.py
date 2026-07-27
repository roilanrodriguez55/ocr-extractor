"""Tests for the DOCX reader."""

import pytest
from docx import Document

from ocr_extractor.readers.docx import read_docx


# ---------- synthetic DOCX fixtures ----------------------------------------


def _make_docx(path, paragraphs=None, table_rows=None):
    """Build a DOCX with optional paragraphs and a single table.

    Parameters
    ----------
    path : pathlib.Path
        Where to save the .docx.
    paragraphs : list[str]
        Plain paragraphs to add in order.
    table_rows : list[list[str]] or None
        2D table contents; first row treated as headers by python-docx.
    """
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text
    doc.save(str(path))
    return path


# ---------- tests ----------------------------------------------------------


class TestReadDocx:
    def test_empty_doc_returns_empty_string(self, tmp_workdir):
        path = _make_docx(tmp_workdir / "empty.docx")
        text = read_docx(path, verbose=False)
        assert text == ""

    def test_single_paragraph(self, tmp_workdir):
        path = _make_docx(
            tmp_workdir / "single.docx",
            paragraphs=["Hello world"],
        )
        text = read_docx(path, verbose=False)
        assert text == "Hello world"

    def test_multiple_paragraphs_joined_by_newlines(self, tmp_workdir):
        path = _make_docx(
            tmp_workdir / "multi.docx",
            paragraphs=["First paragraph", "Second paragraph", "Third paragraph"],
        )
        text = read_docx(path, verbose=False)
        assert text == "First paragraph\nSecond paragraph\nThird paragraph"

    def test_drops_empty_paragraphs(self, tmp_workdir):
        path = _make_docx(
            tmp_workdir / "with_empty.docx",
            paragraphs=["Real content", "", "   ", "More content"],
        )
        text = read_docx(path, verbose=False)
        # Only the non-empty, non-whitespace paragraphs survive.
        assert text == "Real content\nMore content"

    def test_collapses_internal_whitespace(self, tmp_workdir):
        # DOCX paragraphs may contain runs with awkward whitespace.
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello")
        para.add_run("   ")
        para.add_run("world")
        doc.save(str(tmp_workdir / "spaces.docx"))

        text = read_docx(tmp_workdir / "spaces.docx", verbose=False)
        assert text == "Hello world"

    def test_preserves_punctuation(self, tmp_workdir):
        """Unlike the OCR clean_line, DOCX extraction keeps commas and
        other punctuation because the source is already clean.
        """
        path = _make_docx(
            tmp_workdir / "punct.docx",
            paragraphs=["Hello, world! How are you?"],
        )
        text = read_docx(path, verbose=False)
        # Commas and other punctuation are preserved.
        assert "," in text
        assert "!" in text
        assert "?" in text

    def test_no_page_markers(self, tmp_workdir):
        # DOCX has no clean "page" concept, so no PAGE markers are emitted.
        path = _make_docx(
            tmp_workdir / "no_pages.docx",
            paragraphs=["Some content"],
        )
        text = read_docx(path, verbose=False)
        assert "=== PAGE" not in text
        assert "=== END PAGE" not in text

    def test_extracts_table_rows(self, tmp_workdir):
        path = _make_docx(
            tmp_workdir / "with_table.docx",
            paragraphs=["Above the table"],
            table_rows=[["A", "B", "C"], ["1", "2", "3"]],
        )
        text = read_docx(path, verbose=False)
        assert "Above the table" in text
        # Cells joined with " | " separator; leading/trailing pipes
        # from empty edge cells are stripped.
        assert "A | B | C" in text
        assert "1 | 2 | 3" in text

    def test_table_paragraph_order(self, tmp_workdir):
        """A table interleaved between paragraphs keeps document order."""
        doc = Document()
        doc.add_paragraph("Para before")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "X"
        table.rows[0].cells[1].text = "Y"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        doc.add_paragraph("Para after")
        doc.save(str(tmp_workdir / "order.docx"))

        text = read_docx(tmp_workdir / "order.docx", verbose=False)
        before_idx = text.index("Para before")
        table_idx = text.index("X")
        after_idx = text.index("Para after")
        assert before_idx < table_idx < after_idx

    def test_drops_completely_empty_table_rows(self, tmp_workdir):
        path = _make_docx(
            tmp_workdir / "empty_rows.docx",
            table_rows=[["A", "B"], ["", ""], ["", "C"]],
        )
        text = read_docx(path, verbose=False)
        # Row ["A", "B"] survives.
        assert "A" in text
        assert "B" in text
        # Row ["", ""] is entirely empty → dropped (no extra newlines).
        assert text.count("\n") == 1  # only one surviving row
        # Row ["", "C"] has one cell → survives.
        assert "C" in text
        # Leading/trailing pipes are stripped from each row.
        assert not any(line.startswith(" |") for line in text.split("\n"))

    def test_verbose_prints_progress(self, tmp_workdir, capsys):
        path = _make_docx(tmp_workdir / "v.docx", paragraphs=["x"])
        read_docx(path, verbose=True)
        captured = capsys.readouterr()
        assert "Reading" in captured.out

    def test_missing_file_raises(self, tmp_workdir):
        with pytest.raises(FileNotFoundError):
            read_docx(tmp_workdir / "nope.docx", verbose=False)